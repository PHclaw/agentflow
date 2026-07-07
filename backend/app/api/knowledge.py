"""
知识库 API - 文档管理和向量搜索
"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Depends, Form
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import os
import uuid
import aiofiles
from pathlib import Path
from datetime import datetime

from ..core.config import settings
from ..core.database import get_db
from ..models.agent import KnowledgeBase
from ..models.document import Document, DocumentChunk
from ..services.knowledge import KnowledgeService

router = APIRouter(prefix="/api/v1/knowledge", tags=["知识库"])

UPLOAD_DIR = Path(settings.UPLOAD_DIR)
UPLOAD_DIR.mkdir(exist_ok=True)


class CreateKBRequest(BaseModel):
    name: str
    description: str = ""


class SearchRequest(BaseModel):
    query: str
    kb_ids: Optional[List[str]] = None
    top_k: int = 5


# ===== 知识库管理 =====

@router.get("")
async def list_knowledge_bases(db: AsyncSession = Depends(get_db)):
    """列出所有知识库"""
    from sqlalchemy import func

    stmt = select(KnowledgeBase).order_by(KnowledgeBase.created_at.desc())
    result = await db.execute(stmt)
    kbs = result.scalars().all()

    response = []
    for kb in kbs:
        # 获取文档数量
        doc_count = await db.execute(
            select(func.count(Document.id)).where(Document.kb_id == kb.id)
        )
        doc_count = doc_count.scalar() or 0

        response.append({
            "id": kb.id,
            "name": kb.name,
            "description": kb.description,
            "document_count": doc_count,
            "chunk_count": 0,  # 暂时返回 0，后续可以优化查询
            "created_at": kb.created_at.isoformat() if kb.created_at else None,
            "status": "ready",
        })

    return response


@router.post("")
async def create_knowledge_base(
    data: CreateKBRequest,
    db: AsyncSession = Depends(get_db)
):
    """创建知识库"""
    kb = KnowledgeBase(
        id=str(uuid.uuid4()),
        name=data.name,
        description=data.description,
    )
    db.add(kb)
    await db.commit()
    await db.refresh(kb)

    return {
        "id": kb.id,
        "name": kb.name,
        "description": kb.description,
        "message": "创建成功"
    }


@router.get("/{kb_id}")
async def get_knowledge_base(kb_id: str, db: AsyncSession = Depends(get_db)):
    """获取知识库详情"""
    stmt = select(KnowledgeBase).where(KnowledgeBase.id == kb_id)
    result = await db.execute(stmt)
    kb = result.scalar_one_or_none()

    if not kb:
        raise HTTPException(404, "知识库不存在")

    # 获取文档数量
    doc_count = await db.execute(
        select(Document).where(Document.kb_id == kb_id)
    )
    docs = doc_count.scalars().all()

    return {
        "id": kb.id,
        "name": kb.name,
        "description": kb.description,
        "document_count": len(docs),
        "documents": [{
            "id": doc.id,
            "filename": doc.filename,
            "file_type": doc.file_type,
            "status": doc.status,
            "chunk_count": doc.chunk_count,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
        } for doc in docs],
    }


@router.delete("/{kb_id}")
async def delete_knowledge_base(kb_id: str, db: AsyncSession = Depends(get_db)):
    """删除知识库"""
    from sqlalchemy import delete

    # 先删除文档
    await db.execute(delete(Document).where(Document.kb_id == kb_id))
    # 删除知识库
    await db.execute(delete(KnowledgeBase).where(KnowledgeBase.id == kb_id))
    await db.commit()

    return {"message": "删除成功"}


# ===== 文档管理 =====

@router.post("/{kb_id}/documents")
async def upload_document(
    kb_id: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db)
):
    """上传文档到知识库"""
    # 验证知识库存在
    kb = await db.get(KnowledgeBase, kb_id)
    if not kb:
        raise HTTPException(404, "知识库不存在")

    # 验证文件类型
    file_ext = Path(file.filename).suffix.lower()
    allowed_types = {".txt", ".md", ".pdf", ".doc", ".docx"}

    if file_ext not in allowed_types:
        raise HTTPException(400, f"不支持的文件类型：{file_ext}，仅支持 {', '.join(allowed_types)}")

    # 检查文件大小
    content = await file.read()
    if len(content) > settings.MAX_FILE_SIZE:
        raise HTTPException(400, f"文件过大，最大支持 {settings.MAX_FILE_SIZE // 1024 // 1024}MB")

    # 保存文件
    file_id = str(uuid.uuid4())
    file_path = UPLOAD_DIR / "documents" / kb_id / f"{file_id}{file_ext}"
    file_path.parent.mkdir(parents=True, exist_ok=True)

    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(content)

    # 创建文档记录
    doc = Document(
        kb_id=kb_id,
        filename=file.filename,
        file_type=file_ext[1:],  # 去掉点
        file_size=len(content),
        status="processing",
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    try:
        # 提取文本并处理
        text_content = await extract_text(content, file_ext)

        service = KnowledgeService(db, kb_id)
        await service.process_document(
            doc_id=doc.id,
            content=text_content,
            file_path=str(file_path),
        )

        return {
            "id": doc.id,
            "filename": doc.filename,
            "file_type": doc.file_type,
            "status": "completed",
            "message": "上传成功，已处理完成"
        }

    except Exception as e:
        doc.status = "failed"
        await db.commit()
        raise HTTPException(500, f"文档处理失败：{str(e)}")


async def extract_text(content: bytes, file_ext: str) -> str:
    """提取文件文本内容"""
    if file_ext in [".txt", ".md"]:
        return content.decode("utf-8", errors="ignore")

    elif file_ext == ".pdf":
        try:
            import io
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(content))
            text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
            return "\n\n".join(text_parts)
        except Exception:
            return content.decode("utf-8", errors="ignore")

    elif file_ext == ".docx":
        try:
            import io
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception:
            return content.decode("utf-8", errors="ignore")

    return content.decode("utf-8", errors="ignore")


@router.get("/{kb_id}/documents")
async def list_documents(kb_id: str, db: AsyncSession = Depends(get_db)):
    """列出知识库中的文档"""
    stmt = select(Document).where(Document.kb_id == kb_id).order_by(Document.created_at.desc())
    result = await db.execute(stmt)
    docs = result.scalars().all()

    return [{
        "id": doc.id,
        "filename": doc.filename,
        "file_type": doc.file_type,
        "file_size": doc.file_size,
        "chunk_count": doc.chunk_count,
        "status": doc.status,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
    } for doc in docs]


@router.delete("/{kb_id}/documents/{doc_id}")
async def delete_document(kb_id: str, doc_id: int, db: AsyncSession = Depends(get_db)):
    """删除文档"""
    doc = await db.get(Document, doc_id)
    if not doc or doc.kb_id != kb_id:
        raise HTTPException(404, "文档不存在")

    # 删除文档和分块
    from sqlalchemy import delete
    await db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == doc_id))
    await db.delete(doc)
    await db.commit()

    return {"message": "删除成功"}


# ===== 向量搜索 =====

@router.post("/search")
async def search_knowledge(data: SearchRequest, db: AsyncSession = Depends(get_db)):
    """语义搜索"""
    if not data.kb_ids:
        # 获取所有知识库
        stmt = select(KnowledgeBase.id)
        result = await db.execute(stmt)
        data.kb_ids = [row[0] for row in result.fetchall()]

    service = KnowledgeService(db)
    results = await service.search(
        query=data.query,
        kb_ids=data.kb_ids,
        top_k=data.top_k,
    )

    return {
        "query": data.query,
        "results": results,
    }


@router.post("/{kb_id}/search")
async def search_kb(
    kb_id: str,
    data: SearchRequest,
    db: AsyncSession = Depends(get_db)
):
    """在指定知识库中搜索"""
    service = KnowledgeService(db, kb_id)
    results = await service.search_in_kb(
        query=data.query,
        top_k=data.top_k,
    )

    return {
        "kb_id": kb_id,
        "query": data.query,
        "results": results,
    }
