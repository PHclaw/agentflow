"""PDF → Word（.docx；可选再转旧版 .doc）。"""
from __future__ import annotations

import re
import shutil
import subprocess
import uuid
from pathlib import Path
from typing import Any

from app.logging_setup import get_logger

logger = get_logger("pdf-word")

ROOT = Path(__file__).resolve().parents[4]
GENERATED_DIR = ROOT / "static" / "generated"

_IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def _public_url(filename: str) -> str:
    return f"/static/generated/{filename}"


def _safe_stem(name: str) -> str:
    stem = Path(name or "document").stem
    stem = re.sub(r"[^\w.\-一-龥]+", "_", stem) or "document"
    return stem[:80]


def _local_from_url(url: str) -> Path | None:
    u = (url or "").strip()
    if not u:
        return None
    if u.startswith("/static/generated/"):
        p = ROOT / u.lstrip("/")
        return p if p.is_file() else None
    if u.startswith("static/generated/"):
        p = ROOT / u
        return p if p.is_file() else None
    raw = Path(u)
    return raw if raw.is_file() else None


def _pdf_text_fallback(pdfs: list[dict]) -> str:
    import pdfplumber

    chunks: list[str] = []
    for src in pdfs:
        path = Path(str(src.get("path") or ""))
        if not path.is_file():
            continue
        title = src.get("name") or path.name
        chunks.append(f"# {title}")
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    text = (page.extract_text() or "").strip()
                    chunks.append(f"## 第 {i} 页")
                    chunks.append(text or "（本页无抽取文本）")
        except Exception as exc:  # noqa: BLE001
            logger.warning("pdfplumber fallback fail path=%s err=%s", path, exc)
            chunks.append(f"（文本抽取失败：{exc}）")
    return "\n\n".join(chunks).strip()


def _add_runs(paragraph, text: str) -> None:
    """极简：粗体 **x** 拆成 run。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    if len(parts) == 1:
        paragraph.add_run(text)
        return
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def markdown_to_docx(md: str, out_path: Path, *, image_urls: list[str] | None = None) -> None:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.name = "Calibri"
    try:
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    except Exception:  # noqa: BLE001
        pass

    used_images: set[str] = set()
    lines = (md or "").replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        if not line.strip():
            i += 1
            continue
        hm = _HEADING_RE.match(line)
        if hm:
            level = min(len(hm.group(1)), 4)
            doc.add_heading(hm.group(2).strip() or " ", level=level)
            i += 1
            continue
        if line.lstrip().startswith("|") and "|" in line[1:]:
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if cells and re.match(r"^:?-{2,}:?$", cells[0].replace(" ", "")):
                    i += 1
                    continue
                if not all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        table.cell(ri, ci).text = row[ci] if ci < len(row) else ""
            continue
        img_m = _IMG_RE.search(line)
        if img_m:
            caption = img_m.group(1).strip()
            loc = _local_from_url(img_m.group(2).strip())
            if loc:
                try:
                    doc.add_picture(str(loc), width=Inches(5.5))
                    used_images.add(str(loc.resolve()))
                except Exception as exc:  # noqa: BLE001
                    logger.warning("embed image fail %s %s", loc, exc)
            if caption:
                p = doc.add_paragraph(caption)
                p.italic = True
            rest = _IMG_RE.sub("", line).strip()
            if rest:
                _add_runs(doc.add_paragraph(), rest)
            i += 1
            continue
        bullet = re.match(r"^[-*+]\s+(.*)$", line)
        if bullet:
            _add_runs(doc.add_paragraph(style="List Bullet"), bullet.group(1))
            i += 1
            continue
        numbered = re.match(r"^\d+[.)、]\s+(.*)$", line)
        if numbered:
            _add_runs(doc.add_paragraph(style="List Number"), numbered.group(1))
            i += 1
            continue
        _add_runs(doc.add_paragraph(), line)
        i += 1

    for url in image_urls or []:
        loc = _local_from_url(str(url))
        if not loc:
            continue
        key = str(loc.resolve())
        if key in used_images:
            continue
        try:
            doc.add_picture(str(loc), width=Inches(5.5))
            used_images.add(key)
        except Exception as exc:  # noqa: BLE001
            logger.warning("append image fail %s %s", loc, exc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _soffice_bin() -> str | None:
    found = shutil.which("soffice") or shutil.which("soffice.exe")
    if found:
        return found
    for p in (
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        Path("/usr/bin/soffice"),
        Path("/usr/bin/libreoffice"),
    ):
        if p.is_file():
            return str(p)
    return None


def try_docx_to_doc(docx_path: Path) -> Path | None:
    """尽力把 docx 转成旧版 .doc；失败返回 None。"""
    out_dir = docx_path.parent
    target = out_dir / (docx_path.stem + ".doc")
    soffice = _soffice_bin()
    if soffice:
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "doc",
                    "--outdir",
                    str(out_dir),
                    str(docx_path),
                ],
                check=True,
                timeout=90,
                capture_output=True,
            )
            if target.is_file():
                return target
        except Exception as exc:  # noqa: BLE001
            logger.warning("soffice doc convert fail: %s", exc)
    try:
        import win32com.client  # type: ignore
    except Exception:
        return None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        # 0 = wdFormatDocument
        doc.SaveAs(str(target.resolve()), FileFormat=0)
        doc.Close(False)
        word.Quit()
        if target.is_file():
            return target
    except Exception as exc:  # noqa: BLE001
        logger.warning("Word COM doc convert fail: %s", exc)
    return None


def convert_pdfs_to_word(pdfs: list[dict], *, fmt: str = "docx") -> dict[str, Any]:
    """fmt: docx | doc。用户说 word 时应传入 docx。"""
    want_doc = fmt == "doc"
    from app.services.skill.doc_parser import parse_uploaded

    parsed = parse_uploaded(pdfs, intent="to_docx")
    md = str(parsed.get("markdown") or "").strip()
    if not md:
        md = _pdf_text_fallback(pdfs)
    if not md.strip():
        return {
            "intent": "to_doc" if want_doc else "to_docx",
            "script": "pdf_word",
            "exitCode": 1,
            "stdout": parsed.get("stdout") or "",
            "stderr": parsed.get("stderr") or "",
            "note": "未能从 PDF 抽出可用于 Word 的文本。扫描件可能需要 OCR。",
        }

    stem = _safe_stem(str(pdfs[0].get("name") or "document"))
    uid = uuid.uuid4().hex[:8]
    docx_name = f"{stem}-{uid}.docx"
    docx_path = GENERATED_DIR / docx_name
    image_urls = [str(u) for u in (parsed.get("downloadUrls") or [])]
    try:
        markdown_to_docx(md, docx_path, image_urls=image_urls)
    except Exception as exc:  # noqa: BLE001
        logger.error("markdown_to_docx fail: %s", exc)
        return {
            "intent": "to_doc" if want_doc else "to_docx",
            "script": "pdf_word",
            "exitCode": 1,
            "stdout": "",
            "stderr": str(exc),
            "note": f"生成 Word 失败: {exc}",
        }

    if want_doc:
        doc_path = try_docx_to_doc(docx_path)
        if doc_path and doc_path.is_file():
            return {
                "intent": "to_doc",
                "script": "pdf_word",
                "exitCode": 0,
                "stdout": f"format=doc\ndownloadUrl={_public_url(doc_path.name)}\n",
                "stderr": "",
                "note": "已将 PDF 转为旧版 Word（.doc）",
                "downloadUrl": _public_url(doc_path.name),
                "downloadName": doc_path.name,
                "wordFormat": "doc",
            }
        return {
            "intent": "to_doc",
            "script": "pdf_word",
            "exitCode": 0,
            "stdout": f"format=docx (doc fallback)\ndownloadUrl={_public_url(docx_name)}\n",
            "stderr": "",
            "note": (
                "已生成 .docx。本机未安装 LibreOffice 或 Microsoft Word，无法再另存为旧版 .doc。"
                "请用 Word 打开后「另存为」.doc。"
            ),
            "downloadUrl": _public_url(docx_name),
            "downloadName": docx_name,
            "wordFormat": "docx",
        }

    return {
        "intent": "to_docx",
        "script": "pdf_word",
        "exitCode": 0,
        "stdout": f"format=docx\ndownloadUrl={_public_url(docx_name)}\n",
        "stderr": "",
        "note": "已将 PDF 转为 Word（.docx）。用户说 word 时默认即为 docx。",
        "downloadUrl": _public_url(docx_name),
        "downloadName": docx_name,
        "wordFormat": "docx",
    }
